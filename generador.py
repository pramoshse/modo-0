import streamlit as st
import subprocess
import os

# 1. Instalación inmediata de Playwright
@st.cache_resource
def install_playwright():
    # Usamos el comando simple que funciona en el otro código
    subprocess.run(["playwright", "install", "chromium"])

install_playwright()

# Ahora sí, el resto de imports
import asyncio
from playwright.async_api import async_playwright
import datetime
import base64
# ... resto de tus imports (openpyxl, PIL, docx, etc.)

import html
import io
import json
import re
from typing import Any, Dict, Iterable, List, Optional

import streamlit.components.v1 as components


# ============================================================
# CONFIGURACIÓN GENERAL
# ============================================================

st.set_page_config(
    page_title="Generador de Procedimientos Modo 0",
    layout="wide",
)

st.markdown(
    """
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800;900&display=swap');

        html, body, [class*="css"], .stApp {
            font-family: 'Inter', Arial, sans-serif;
        }

        .main .block-container {
            padding-top: 1.6rem;
            max-width: 1320px;
        }

        .soft-note {
            font-size: 0.78rem;
            color: #64748B;
            line-height: 1.45;
            margin-top: -0.35rem;
            margin-bottom: 1rem;
        }

        .status-ok {
            border: 1px solid #BBF7D0;
            background: #F0FDF4;
            color: #166534;
            border-radius: 10px;
            padding: 0.8rem 1rem;
            font-size: 0.9rem;
        }

        .status-warn {
            border: 1px solid #FDE68A;
            background: #FFFBEB;
            color: #92400E;
            border-radius: 10px;
            padding: 0.8rem 1rem;
            font-size: 0.9rem;
        }
    </style>
    """,
    unsafe_allow_html=True,
)


# ============================================================
# UTILIDADES DE DATOS
# ============================================================


def _html(value: Any) -> str:
    """Escapa texto para insertarlo de forma segura en HTML."""
    if value is None:
        return ""
    return html.escape(str(value), quote=True)



def _normalize(value: Any) -> str:
    if value is None:
        return ""
    value = str(value).strip()
    value = re.sub(r"\s+", " ", value)
    return value



def _upper(value: Any) -> str:
    return _normalize(value).upper()



def _get_from_dict(source: Dict[str, Any], keys: Iterable[str], default: str = "") -> Any:
    """Busca el primer valor no vacío dentro de un diccionario usando nombres alternativos."""
    for key in keys:
        if key in source and source.get(key) not in (None, "", [], {}):
            return source.get(key)
    return default



def _get_nested(source: Dict[str, Any], path: Iterable[str], default: Any = "") -> Any:
    current: Any = source
    for key in path:
        if not isinstance(current, dict) or key not in current:
            return default
        current = current[key]
    return current if current not in (None, "") else default



def _split_tasks(raw: Any) -> List[str]:
    """Convierte tareas en lista limpia, aceptando texto libre, listas, tuplas o diccionarios."""
    if raw is None:
        return []

    if isinstance(raw, dict):
        for key in ("tareas", "items", "seleccionadas", "predefinidas", "values"):
            if key in raw and raw.get(key) not in (None, "", [], {}):
                return _split_tasks(raw.get(key))
        candidate = _get_from_dict(
            raw,
            ["tarea", "actividad", "descripcion", "descripción", "descripcion_tarea", "nombre", "texto", "label"],
            "",
        )
        return [str(candidate).strip(" -•	")] if str(candidate).strip() else []

    if isinstance(raw, (list, tuple)):
        tasks: List[str] = []
        for item in raw:
            if isinstance(item, dict):
                tasks.extend(_split_tasks(item))
            else:
                item_text = str(item).strip(" -•	")
                if item_text:
                    tasks.append(item_text)
        return [task for task in tasks if task]

    text = str(raw).strip()
    if not text:
        return []

    parts = []
    for line in re.split(r"[\n\r]+", text):
        line = line.strip(" -•	")
        if not line:
            continue
        # Si el usuario cargó una frase larga separada por comas, se conserva como una tarea completa.
        parts.append(line)
    return parts



def _mode_is_modo_0(mode: Any) -> bool:
    clean = _normalize(mode).lower().replace("°", "")
    return clean in {"modo 0", "modo0", "0", "m0"}



def _auto_negocio_from_sitio(sitio: Any, default: str = "") -> str:
    """Determina el negocio por regla solicitada: Planta/CEDI -> Bebidas; resto -> Ingenios."""
    sitio_txt = _normalize(sitio).lower()
    if not sitio_txt:
        return _normalize(default)
    if sitio_txt.startswith("planta") or sitio_txt.startswith("cedi"):
        return "Bebidas"
    return "Ingenios"


def _bytes_to_data_uri(content: bytes, mime: str = "image/png") -> str:
    if not content:
        return ""
    encoded = base64.b64encode(content).decode("utf-8")
    return f"data:{mime or 'image/png'};base64,{encoded}"


def _mime_from_filename(filename: str, default: str = "image/png") -> str:
    ext = os.path.splitext(str(filename or ""))[1].lower()
    return {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".webp": "image/webp",
    }.get(ext, default)


def _uploaded_file_data_uri(uploaded_file) -> str:
    if uploaded_file is None:
        return ""
    try:
        content = uploaded_file.getvalue()
        mime = getattr(uploaded_file, "type", "") or _mime_from_filename(getattr(uploaded_file, "name", ""))
        return _bytes_to_data_uri(content, mime)
    except Exception:
        return ""


def _local_file_data_uri(candidates: Iterable[str]) -> str:
    for candidate in candidates:
        if not candidate:
            continue
        try:
            if os.path.exists(candidate) and os.path.isfile(candidate):
                with open(candidate, "rb") as fh:
                    return _bytes_to_data_uri(fh.read(), _mime_from_filename(candidate))
        except Exception:
            continue
    return ""



def _photo_data_uri(payload: Dict[str, Any]) -> str:
    photo = payload.get("equipment_photo") or {}
    content = photo.get("content_base64")
    mime = photo.get("mime_type") or "image/png"
    if not content:
        return ""
    return f"data:{mime};base64,{content}"



def _logo_data_uri(uploaded_logo=None) -> str:
    """Carga el logo local arca.png por defecto y permite reemplazarlo por archivo subido."""
    uploaded_uri = _uploaded_file_data_uri(uploaded_logo)
    if uploaded_uri:
        return uploaded_uri

    base_dir = os.path.dirname(os.path.abspath(__file__)) if "__file__" in globals() else os.getcwd()
    return _local_file_data_uri(
        [
            os.path.join(base_dir, "arca.png"),
            os.path.join(os.getcwd(), "arca.png"),
            "/home/ubuntu/arca.png",
            "/home/ubuntu/upload/arca.png",
        ]
    )



def extract_procedure_context(payload: Dict[str, Any]) -> Dict[str, Any]:
    """Extrae datos del borrador de forma tolerante a cambios de nombre de variables."""
    widgets = payload.get("widgets") if isinstance(payload.get("widgets"), dict) else {}
    computed = payload.get("computed") if isinstance(payload.get("computed"), dict) else {}

    mode_final = _get_from_dict(
        computed,
        ["modo_final", "modo_convalidado", "modo_resultante"],
        _get_from_dict(widgets, ["modo_final", "modo_convalidado", "modo_resultante"], ""),
    )

    tasks_raw = _get_from_dict(
        widgets,
        [
            "tareas_predefinidas",
            "tareas_predefinidas_txt",
            "tareas_preseleccionadas",
            "tareas_seleccionadas",
        ],
        "",
    )
    if not tasks_raw:
        tasks_raw = _get_from_dict(
            payload,
            [
                "tareas_predefinidas",
                "tareas_predefinidas_txt",
                "tareas_preseleccionadas",
                "tareas_seleccionadas",
            ],
            "",
        )
    if not tasks_raw:
        tasks_raw = _get_from_dict(
            widgets,
            [
                "tareas_txt",
                "tareas",
                "tareas_aplicables",
                "descripcion_tareas",
                "actividad",
                "actividades",
            ],
            "",
        )
    tasks = _split_tasks(tasks_raw)
    if not tasks:
        tasks = [
            "Operación normal del equipo con protecciones instaladas y funcionales.",
            "Interacción mediante panel HMI o controles normales del operador.",
            "No se anulan, retiran ni desvían resguardos o dispositivos de seguridad.",
        ]

    energias = _get_from_dict(computed, ["energias_seleccionadas", "energias"], {})
    lista_peligros = _get_from_dict(computed, ["lista_peligros_final", "peligros"], [])
    evaluaciones_fine = _get_from_dict(computed, ["evaluaciones_fine"], [])

    sitio = _get_from_dict(widgets, ["sitio", "planta", "site"], "")
    negocio_json = _get_from_dict(widgets, ["negocio", "business", "unidad_negocio"], "")

    return {
        "negocio": _auto_negocio_from_sitio(sitio, negocio_json),
        "sitio": sitio,
        "tipo_sitio": _get_from_dict(widgets, ["tipo_sitio"], ""),
        "area": _get_from_dict(widgets, ["area_sector", "area", "sector", "área"], ""),
        "linea": _get_from_dict(widgets, ["linea", "línea", "linea_equipo"], ""),
        "equipo": _get_from_dict(
            widgets,
            ["equipo_desc", "equipo", "descripcion_equipo", "nombre_equipo", "maquina", "máquina"],
            "EQUIPO SIN DESCRIPCIÓN",
        ),
        "fabricante": _get_from_dict(widgets, ["fabricante", "marca"], ""),
        "modelo": _get_from_dict(widgets, ["modelo"], ""),
        "anio": _get_from_dict(widgets, ["anio", "año"], ""),
        "modo_inicial": _get_from_dict(computed, ["modo_inicial"], _get_from_dict(widgets, ["modo_inicial"], "")),
        "modo_final": mode_final,
        "tareas": tasks,
        "energias": energias,
        "lista_peligros": lista_peligros,
        "evaluaciones_fine": evaluaciones_fine,
        "photo_uri": "",
    }


# ============================================================
# PLANTILLA HTML/PDF MODO 0
# ============================================================


def _tasks_html(tasks: List[str]) -> str:
    return "".join(f"<div class='task-line'>- {_html(task)}</div>" for task in tasks)



def _equipment_meta(ctx: Dict[str, Any]) -> str:
    meta = []
    if ctx.get("fabricante"):
        meta.append(f"Fabricante: {_html(ctx.get('fabricante'))}")
    if ctx.get("modelo"):
        meta.append(f"Modelo: {_html(ctx.get('modelo'))}")
    if ctx.get("anio"):
        meta.append(f"Año: {_html(ctx.get('anio'))}")
    return " · ".join(meta)



def build_modo_0_html(
    ctx: Dict[str, Any],
    *,
    codigo: str,
    revision: str,
    fecha: datetime.date,
    organizacion: str,
    logo_uri: str = "",
    personal_afectado: str,
    personal_autorizado: str,
    elaborado_por: str,
    aprobado_por: str,
    puesto_elaborado: str = "",
    puesto_aprobado: str = "",
    fecha_firma: datetime.date | str | None = None,
) -> str:
    fecha_txt = fecha.strftime("%d/%m/%Y") if isinstance(fecha, datetime.date) else _normalize(fecha)
    fecha_firma_txt = fecha_firma.strftime("%d/%m/%Y") if isinstance(fecha_firma, datetime.date) else (_normalize(fecha_firma) or fecha_txt)
    task_rows = _tasks_html(ctx.get("tareas") or [])
    photo_uri = ctx.get("photo_uri", "")
    equipo_meta = _equipment_meta(ctx)

    logo_block = (
        f"<img class='logo-img' src='{logo_uri}' alt='Logo'>"
        if logo_uri
        else f"<div class='logo-text'>{_html(organizacion)}</div>"
    )

    photo_block = (
        f"<img class='equipment-photo' src='{photo_uri}' alt='Foto del paso a paso'>"
        if photo_uri
        else "<div class='photo-placeholder'>FOTO DEL PASO A PASO</div>"
    )

    html_doc = f"""
<!doctype html>
<html lang="es">
<head>
<meta charset="utf-8">
<title>Procedimiento Modo 0 - Control de Energías Peligrosas</title>
<style>
    @page {{
        size: A4 portrait;
        margin: 5mm;
    }}

    * {{
        box-sizing: border-box;
    }}

    body {{
        margin: 0;
        background: #E5E7EB;
        color: #0F172A;
        font-family: Bahnschrift, 'Bahnschrift SemiCondensed', 'Arial Narrow', Arial, Helvetica, sans-serif;
        font-size: 8.7px;
        line-height: 1.14;
    }}

    .sheet {{
        width: 760px;
        min-height: 0;
        margin: 0 auto;
        background: #FFFFFF;
        border: 1.2px solid #111827;
        box-shadow: 0 12px 28px rgba(15, 23, 42, 0.18);
    }}

    table {{
        border-collapse: collapse;
        width: 100%;
        table-layout: fixed;
    }}

    td, th {{
        border: 1px solid #111827;
        padding: 5px 6px;
        vertical-align: middle;
    }}

    .top-black {{
        background: #050505;
        color: #FFFFFF;
        font-weight: 800;
        text-align: center;
        letter-spacing: 0.06em;
        font-size: 11px;
        padding: 2px 4px;
        text-transform: uppercase;
    }}

    .logo-cell {{
        width: 62px;
        text-align: center;
        background: #FFFFFF;
        padding: 3px;
    }}

    .logo-img {{
        max-width: 56px;
        max-height: 40px;
        object-fit: contain;
    }}

    .logo-text {{
        color: #B91C1C;
        font-size: 9px;
        font-weight: 900;
        line-height: 1.05;
        text-transform: uppercase;
        word-break: break-word;
    }}

    .main-title {{
        text-align: center;
        font-size: 12.5px;
        font-weight: 900;
        letter-spacing: 0.01em;
        text-transform: uppercase;
        padding: 5px 6px;
    }}

    .doc-data-label {{
        background: #E5E7EB;
        width: 68px;
        text-align: center;
        font-weight: 800;
        text-transform: uppercase;
        font-size: 10px;
    }}

    .doc-data-value {{
        width: 86px;
        text-align: center;
        font-weight: 700;
        background: #F8FAFC;
        font-size: 10px;
    }}

    .info-label {{
        display: block;
        color: #334155;
        font-size: 10px;
        margin-bottom: 2px;
    }}

    .info-value {{
        display: block;
        color: #111827;
        font-size: 13px;
        font-weight: 700;
    }}

    .info-box {{
        height: 50px;
        text-align: center;
        background: #FFFFFF;
    }}

    .person-title {{
        font-weight: 900;
        text-align: center;
        background: #69C97F;
        color: #0F172A;
        font-size: 8.8px;
        padding: 3px 6px;
    }}

    .person-body {{
        text-align: center;
        min-height: 28px;
        font-size: 8.8px;
        white-space: pre-line;
    }}

    .equipment-label {{
        width: 62px;
        font-weight: 800;
        color: #334155;
        background: #F8FAFC;
        text-align: center;
    }}

    .equipment-title {{
        text-align: center;
        font-size: 15px;
        font-weight: 900;
        text-transform: uppercase;
        padding: 4px 6px 2px 6px;
    }}

    .equipment-meta {{
        display: block;
        margin-top: 3px;
        color: #64748B;
        font-size: 9px;
        font-weight: 600;
        text-transform: none;
    }}

    .red-header th,
    .red-bar {{
        background: #69C97F;
        color: #0F172A;
        font-weight: 900;
        text-align: center;
        font-size: 10px;
    }}

    .red-bar {{
        border-left: 1px solid #111827;
        border-right: 1px solid #111827;
        padding: 4px 8px;
        text-transform: none;
    }}

    .block-zero {{
        text-align: center;
        font-size: 48px;
        font-weight: 900;
        color: #000000;
        background: #FFFFFF;
        line-height: 1;
    }}

    .mode-box {{
        background: #69C97F;
        color: #0F172A;
        text-align: center;
        font-size: 17px;
        font-weight: 900;
        text-transform: uppercase;
    }}

    .tasks-cell {{
        height: 54px;
        font-size: 8.6px;
        background: #FFFFFF;
        padding: 5px 8px;
    }}

    .task-line {{
        margin: 0 0 4px 0;
    }}

    .photo-area {{
        height: 210px;
        background: #FFFFFF;
        text-align: center;
        vertical-align: middle;
        padding: 7px;
    }}

    .equipment-photo {{
        max-width: 100%;
        max-height: 195px;
        object-fit: contain;
        border: 1px solid #CBD5E1;
        background: #F8FAFC;
        padding: 3px;
    }}

    .photo-placeholder {{
        height: 190px;
        display: flex;
        align-items: center;
        justify-content: center;
        color: #94A3B8;
        font-size: 14px;
        font-weight: 800;
        letter-spacing: 0.03em;
        border: 1px dashed #CBD5E1;
        background: #F8FAFC;
    }}

    .dark-head th {{
        background: #3B3B3B;
        color: #FFFFFF;
        font-weight: 900;
        text-align: center;
        font-size: 8.3px;
        padding: 3px 2px;
    }}

    .procedure-table td {{
        height: 40px;
        font-size: 7.8px;
        text-align: center;
        background: #FFFFFF;
    }}

    .procedure-note {{
        text-align: center;
        font-size: 7.2px;
        line-height: 1.12;
        padding: 4px 4px !important;
    }}

    .legend-title {{
        background: #69C97F;
        color: #0F172A;
        font-weight: 900;
        text-align: center;
        padding: 2px;
    }}

    .energy-legend td,
    .lock-legend td {{
        text-align: center;
        font-weight: 900;
        font-size: 7.6px;
        color: #0F172A;
        padding: 3px 2px;
    }}

    .e-electric {{ background: #000000; color: #FFFFFF !important; }}
    .e-neumatic {{ background: #0284C7; color: #FFFFFF !important; }}
    .e-amoniaco {{
        background: linear-gradient(90deg, #D9D9D9 0 16%, #F59E0B 16% 21%, #D9D9D9 21% 43%, #F59E0B 43% 48%, #D9D9D9 48% 70%, #F59E0B 70% 75%, #D9D9D9 75% 100%);
    }}
    .e-termica {{ background: #DC2626; color: #FFFFFF !important; }}
    .e-hidraulica {{ background: #7C3AED; color: #FFFFFF !important; }}
    .e-potencial {{
        background: linear-gradient(90deg, #FFF200 0 22%, #050505 22% 28%, #FFF200 28% 66%, #050505 66% 72%, #FFF200 72% 100%);
    }}
    .e-quimica {{ background: #FFF200; }}
    .e-vapor {{ background: #F59E0B; color: #111827 !important; }}
    .e-agua {{ background: #16A34A; color: #FFFFFF !important; }}
    .e-soda {{
        background: linear-gradient(90deg, #D9D9D9 0 24%, #F59E0B 24% 30%, #D9D9D9 30% 64%, #F59E0B 64% 70%, #D9D9D9 70% 100%);
    }}
    .e-ozono {{ background: #BAE6FD; }}
    .e-gas {{ background: #C7D2FE; color: #111827 !important; }}

    .l-mmto {{ background: #EF0000; color: #FFFFFF !important; }}
    .l-calidad {{ background: #FFF200; }}
    .l-produccion {{ background: #16A34A; color: #FFFFFF !important; }}
    .l-edilicio {{ background: #0F7DBD; color: #FFFFFF !important; }}
    .l-supervisor {{ background: #050505; color: #FFFFFF !important; }}

    .lock-legend {{ table-layout: fixed; }}
    .lock-legend td {{
        width: 20%;
        height: 24px;
        line-height: 1.08;
        vertical-align: middle;
        overflow: hidden;
    }}
    .lock-item {{
        display: inline-flex;
        align-items: center;
        justify-content: center;
        gap: 4px;
        max-width: 100%;
        white-space: normal;
    }}
    .lock-symbol {{
        display: inline-block;
        flex: 0 0 auto;
        position: relative;
        width: 9px;
        height: 7px;
        border: 1.2px solid currentColor;
        border-radius: 1px;
        transform: translateY(2px);
    }}
    .lock-symbol::before {{
        content: "";
        position: absolute;
        left: 1.5px;
        top: -6px;
        width: 4px;
        height: 5px;
        border: 1.2px solid currentColor;
        border-bottom: 0;
        border-radius: 5px 5px 0 0;
    }}

    .footer-sign {{
        margin-top: 0;
    }}

    .signature-cell {{
        height: 30px;
        font-size: 7.8px;
        color: #334155;
        background: #F8FAFC;
        text-align: center;
    }}

    .small-muted {{
        color: #64748B;
        font-size: 8.8px;
        font-weight: 600;
    }}

    @media print {{
        body {{ background: #FFFFFF; }}
        .sheet {{ box-shadow: none; margin: 0; width: 100%; }}
    }}
</style>
</head>
<body>
<div class="sheet">
    <table>
        <tr>
            <td class="logo-cell" rowspan="3">{logo_block}</td>
            <td class="top-black" colspan="2">CONTROL DE ENERGÍAS PELIGROSAS</td>
            <td class="doc-data-label">Código</td>
            <td class="doc-data-value">{_html(codigo)}</td>
        </tr>
        <tr>
            <td class="main-title" colspan="2" rowspan="2">PROCEDIMIENTO ESPECÍFICO PARA CONTROL DE ENERGÍAS</td>
            <td class="doc-data-label">Revisión</td>
            <td class="doc-data-value">{_html(revision)}</td>
        </tr>
        <tr>
            <td class="doc-data-label">Fecha</td>
            <td class="doc-data-value">{_html(fecha_txt)}</td>
        </tr>
    </table>

    <table>
        <tr>
            <td class="info-box" style="width: 62px;">
                <span class="info-label">Negocio:</span>
                <span class="info-value">{_html(ctx.get('negocio') or '-')}</span>
            </td>
            <td class="info-box" style="width: 86px;">
                <span class="info-label">Sitio:</span>
                <span class="info-value">{_html(ctx.get('sitio') or '-')}</span>
            </td>
            <td class="info-box" style="width: 86px;">
                <span class="info-label">Área:</span>
                <span class="info-value">{_html(ctx.get('area') or '-')}</span>
            </td>
            <td class="info-box" style="width: 112px;">
                <span class="info-label">Línea:</span>
                <span class="info-value">{_html(ctx.get('linea') or '-')}</span>
            </td>
            <td style="padding:0;">
                <table>
                    <tr><td class="person-title">Personal afectado - Puestos de trabajo</td></tr>
                    <tr><td class="person-body">{_html(personal_afectado)}</td></tr>
                    <tr><td class="person-title">Personal autorizado - Puestos de trabajo</td></tr>
                    <tr><td class="person-body">{_html(personal_autorizado)}</td></tr>
                </table>
            </td>
        </tr>
    </table>

    <table>
        <tr>
            <td class="equipment-label">Equipo:</td>
            <td class="equipment-title">
                {_html(ctx.get('equipo') or 'EQUIPO')}
                <span class="equipment-meta">{equipo_meta}</span>
            </td>
        </tr>
    </table>

    <table>
        <tr class="red-header">
            <th style="width: 62px;">Puntos de<br>Bloqueo</th>
            <th style="width: 86px;">Modo de<br>Intervención</th>
            <th>Listado de tareas aplicable al presente procedimiento</th>
        </tr>
        <tr>
            <td class="block-zero">0</td>
            <td class="mode-box">MODO 0</td>
            <td class="tasks-cell">{task_rows}</td>
        </tr>
    </table>

    <div class="red-bar">Procedimiento - Control de Energías Peligrosas</div>

    <table>
        <tr>
            <td class="photo-area">{photo_block}</td>
        </tr>
    </table>

    <table class="procedure-table">
        <tr class="dark-head">
            <th style="width: 62px;">Punto de<br>Bloqueo</th>
            <th style="width: 86px;">Fuente de Energía</th>
            <th style="width: 86px;">Magnitud</th>
            <th style="width: 112px;">Ubicación</th>
            <th style="width: 142px;">Acción</th>
            <th style="width: 158px;">Verificación</th>
            <th style="width: 88px;">Dispositivo(s) de<br>Aislamiento</th>
        </tr>
        <tr>
            <td>No aplica</td>
            <td>No Aplica</td>
            <td>No Aplica</td>
            <td>No Aplica</td>
            <td class="procedure-note">
                1. Operar la máquina exclusivamente mediante el panel de Interfaz Hombre-Máquina (HMI) o mediante los controles normales asignados al operador, sin realizar intervención directa sobre el equipo.<br><br>
                2. Mantener la operación dentro de las condiciones normales previstas, sin introducir manos u otros elementos en zonas de riesgo, partes móviles o sectores energizados.<br><br>
                3. No retirar, anular, puentear ni desactivar resguardos, enclavamientos o cualquier dispositivo de seguridad durante la ejecución de la tarea.<br><br>
                4. Detener la actividad y escalar la evaluación si la tarea requiere retirar, superar o intervenir cualquier sistema de protección de la máquina.
            </td>
            <td class="procedure-note">
                1. Confirmar visualmente que todos los resguardos y dispositivos de protección se encuentran instalados, en posición correcta y en condiciones funcionales.<br><br>
                2. Verificar que el equipo opera dentro de sus parámetros normales y que no existe interacción directa del operador con partes móviles, energizadas o zonas de peligro.<br><br>
                3. Comprobar que ninguna protección ha sido modificada, anulada o superada para realizar la actividad.<br><br>
                4. Confirmar que, en caso de requerirse acceso a zonas protegidas o intervención directa sobre el equipo, se haya reevaluado el modo de intervención aplicable, dejando sin efecto el presente procedimiento Modo 0.
            </td>
            <td>No requerido</td>
        </tr>
    </table>

    <div class="legend-title">Clasificación de Energías Peligrosas</div>
    <table class="energy-legend">
        <tr>
            <td class="e-electric">E: Eléctrica</td>
            <td class="e-neumatic">N: Neumática</td>
            <td class="e-amoniaco">AM: Amoníaco</td>
            <td class="e-termica">T: Térmica</td>
            <td class="e-hidraulica">H: Hidráulica</td>
            <td class="e-potencial">P: Potencial</td>
        </tr>
        <tr>
            <td class="e-quimica">Q: Química</td>
            <td class="e-vapor">V: Vapor</td>
            <td class="e-agua">A: Agua</td>
            <td class="e-soda">SC: Soda Cáustica</td>
            <td class="e-ozono">Oz: Ozono</td>
            <td class="e-gas">GC: Gas Carbónico</td>
        </tr>
    </table>

    <div class="legend-title">Clasificación de Candados según sector y función</div>
    <table class="lock-legend">
        <tr>
            <td class="l-mmto"><span class="lock-item"><span class="lock-symbol"></span><span>MMTO<br>Industrial</span></span></td>
            <td class="l-calidad"><span class="lock-item"><span class="lock-symbol"></span><span>Calidad</span></span></td>
            <td class="l-produccion"><span class="lock-item"><span class="lock-symbol"></span><span>Producción</span></span></td>
            <td class="l-edilicio"><span class="lock-item"><span class="lock-symbol"></span><span>Mantenimiento Edilicio<br>Contratistas</span></span></td>
            <td class="l-supervisor"><span class="lock-item"><span class="lock-symbol"></span><span>Supervisor MMTO Industrial<br>(Bloqueo Departamental)</span></span></td>
        </tr>
    </table>

    <table class="footer-sign">
        <tr>
            <td class="signature-cell">
                Elaborado por: <strong>{_html(elaborado_por or '-')}</strong><br>
                <span class="small-muted">Puesto: {_html(puesto_elaborado or '-')} · Fecha: {_html(fecha_firma_txt)}</span>
            </td>
            <td class="signature-cell">
                Aprobado por: <strong>{_html(aprobado_por or '-')}</strong><br>
                <span class="small-muted">Puesto: {_html(puesto_aprobado or '-')} · Fecha: {_html(fecha_firma_txt)}</span>
            </td>
        </tr>
    </table>
</div>
</body>
</html>
"""
    return html_doc



def _get_playwright_sync_api():
    """Carga Playwright de forma diferida para generar PDF desde Chromium."""
    try:
        module = __import__("playwright.sync_api", fromlist=["sync_playwright"])
        return getattr(module, "sync_playwright", None)
    except Exception:
        return None



def html_to_pdf_bytes(html_doc: str) -> bytes:
    """Convierte el HTML del procedimiento en PDF usando Chromium vía Playwright."""
    sync_playwright = _get_playwright_sync_api()
    if sync_playwright is None:
        raise RuntimeError(
            "La exportación PDF requiere Playwright. Instalá la dependencia con: pip install playwright && python -m playwright install chromium."
        )

    try:
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            page = browser.new_page(viewport={"width": 1240, "height": 1754})
            page.set_content(html_doc, wait_until="networkidle")
            page.emulate_media(media="print")
            pdf_bytes = page.pdf(
                format="A4",
                landscape=False,
                scale=0.90,
                print_background=True,
                prefer_css_page_size=True,
                margin={"top": "0mm", "right": "0mm", "bottom": "0mm", "left": "0mm"},
            )
            browser.close()
            return pdf_bytes
    except Exception as exc:
        raise RuntimeError(
            "No se pudo generar el PDF con Playwright. Verificá que Chromium esté instalado con: python -m playwright install chromium. "
            f"Detalle: {exc}"
        ) from exc



def _data_uri_to_bytes(data_uri: str) -> tuple[bytes, str]:
    """Convierte un data URI en bytes y extensión sugerida para insertar imágenes reales."""
    if not data_uri or not isinstance(data_uri, str) or not data_uri.startswith("data:"):
        return b"", ".png"
    header, _, encoded = data_uri.partition(",")
    if not encoded:
        return b"", ".png"
    mime = header.split(";", 1)[0].replace("data:", "").strip().lower()
    ext = {
        "image/png": ".png",
        "image/jpeg": ".jpg",
        "image/jpg": ".jpg",
        "image/webp": ".webp",
    }.get(mime, ".png")
    try:
        return base64.b64decode(encoded), ext
    except Exception:
        return b"", ext


def _fecha_export_txt(value: Any) -> str:
    if isinstance(value, datetime.date):
        return value.strftime("%d/%m/%Y")
    return _normalize(value)


def _docx_set_cell_shading(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.replace("#", "").upper())


def _docx_set_cell_borders(cell, color: str = "111827", size: str = "6") -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color.replace("#", "").upper())


def _docx_set_cell_margins(cell, top: int = 45, start: int = 45, bottom: int = 45, end: int = 45) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def _docx_set_cell_width(cell, width_cm: float) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm

    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def _docx_clear_cell(cell) -> None:
    cell.text = ""


def _docx_write_cell(
    cell,
    text: Any = "",
    *,
    bold: bool = False,
    size: float = 7.5,
    color: str = "0F172A",
    fill: str | None = None,
    align: str = "center",
    valign: str = "center",
) -> None:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    _docx_clear_cell(cell)
    if fill:
        _docx_set_cell_shading(cell, fill)
    _docx_set_cell_borders(cell)
    _docx_set_cell_margins(cell)
    cell.vertical_alignment = {
        "top": WD_CELL_VERTICAL_ALIGNMENT.TOP,
        "bottom": WD_CELL_VERTICAL_ALIGNMENT.BOTTOM,
    }.get(valign, WD_CELL_VERTICAL_ALIGNMENT.CENTER)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
    }.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    paragraph.paragraph_format.space_before = 0
    paragraph.paragraph_format.space_after = 0
    run = paragraph.add_run(str(text or ""))
    run.bold = bold
    run.font.name = "Bahnschrift"
    run.font.size = __import__("docx.shared", fromlist=["Pt"]).Pt(size)
    run.font.color.rgb = __import__("docx.shared", fromlist=["RGBColor"]).RGBColor.from_string(color.replace("#", "").upper())


def _docx_apply_table_grid(table, widths_cm: list[float] | None = None) -> None:
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
    from docx.shared import Cm

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for idx, cell in enumerate(row.cells):
            _docx_set_cell_borders(cell)
            _docx_set_cell_margins(cell)
            if widths_cm and idx < len(widths_cm):
                _docx_set_cell_width(cell, widths_cm[idx])
    if widths_cm:
        for idx, width in enumerate(widths_cm):
            try:
                table.columns[idx].width = Cm(width)
            except Exception:
                pass


def _docx_add_spacer(document, points: float = 1.0) -> None:
    from docx.shared import Pt

    p = document.add_paragraph()
    p.paragraph_format.space_before = 0
    p.paragraph_format.space_after = Pt(points)


def html_to_word_bytes(
    ctx: Dict[str, Any],
    *,
    codigo: str,
    revision: str,
    fecha: datetime.date | str,
    organizacion: str,
    logo_uri: str = "",
    personal_afectado: str,
    personal_autorizado: str,
    elaborado_por: str,
    aprobado_por: str,
    puesto_elaborado: str = "",
    puesto_aprobado: str = "",
    fecha_firma: datetime.date | str | None = None,
) -> bytes:
    """Genera un DOCX editable con tablas, textos, colores y bordes reales."""
    from docx import Document
    from docx.enum.section import WD_SECTION_START
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    fecha_txt = _fecha_export_txt(fecha)
    fecha_firma_txt = _fecha_export_txt(fecha_firma) or fecha_txt
    equipo_meta = _equipment_meta(ctx)

    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(0.45)
    section.bottom_margin = Cm(0.45)
    section.left_margin = Cm(0.45)
    section.right_margin = Cm(0.45)

    style = document.styles["Normal"]
    style.font.name = "Bahnschrift"
    style.font.size = Pt(7.5)

    header = document.add_table(rows=3, cols=5)
    _docx_apply_table_grid(header, [2.0, 5.0, 5.2, 2.4, 3.2])
    header.cell(0, 0).merge(header.cell(2, 0))
    header.cell(0, 1).merge(header.cell(0, 2))
    header.cell(1, 1).merge(header.cell(2, 2))

    logo_cell = header.cell(0, 0)
    _docx_write_cell(logo_cell, "", fill="FFFFFF")
    logo_bytes, _ = _data_uri_to_bytes(logo_uri)
    if logo_bytes:
        try:
            p = logo_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(io.BytesIO(logo_bytes), width=Cm(1.55))
        except Exception:
            _docx_write_cell(logo_cell, organizacion, bold=True, size=7.0, color="B91C1C", fill="FFFFFF")
    else:
        _docx_write_cell(logo_cell, organizacion, bold=True, size=7.0, color="B91C1C", fill="FFFFFF")

    _docx_write_cell(header.cell(0, 1), "CONTROL DE ENERGÍAS PELIGROSAS", bold=True, size=8.5, color="FFFFFF", fill="050505")
    _docx_write_cell(header.cell(1, 1), "PROCEDIMIENTO ESPECÍFICO PARA CONTROL DE ENERGÍAS", bold=True, size=9.0, fill="FFFFFF")
    for row_idx, label, value in ((0, "Código", codigo), (1, "Revisión", revision), (2, "Fecha", fecha_txt)):
        _docx_write_cell(header.cell(row_idx, 3), label, bold=True, size=7.5, fill="E5E7EB")
        _docx_write_cell(header.cell(row_idx, 4), value, bold=True, size=7.5, fill="F8FAFC")

    info = document.add_table(rows=1, cols=5)
    _docx_apply_table_grid(info, [2.0, 2.8, 2.8, 3.6, 6.6])
    for idx, (label, value) in enumerate(
        [
            ("Negocio:", ctx.get("negocio") or "-"),
            ("Sitio:", ctx.get("sitio") or "-"),
            ("Área:", ctx.get("area") or "-"),
            ("Línea:", ctx.get("linea") or "-"),
        ]
    ):
        _docx_write_cell(info.cell(0, idx), f"{label}\n{value}", bold=False, size=7.2, fill="FFFFFF")
    person_cell = info.cell(0, 4)
    _docx_clear_cell(person_cell)
    _docx_set_cell_borders(person_cell)
    nested = person_cell.add_table(rows=4, cols=1)
    _docx_apply_table_grid(nested, [6.4])
    _docx_write_cell(nested.cell(0, 0), "Personal afectado - Puestos de trabajo", bold=True, size=6.8, fill="69C97F")
    _docx_write_cell(nested.cell(1, 0), personal_afectado, size=6.8, fill="FFFFFF")
    _docx_write_cell(nested.cell(2, 0), "Personal autorizado - Puestos de trabajo", bold=True, size=6.8, fill="69C97F")
    _docx_write_cell(nested.cell(3, 0), personal_autorizado, size=6.8, fill="FFFFFF")

    equipment = document.add_table(rows=1, cols=2)
    _docx_apply_table_grid(equipment, [2.0, 15.8])
    _docx_write_cell(equipment.cell(0, 0), "Equipo:", bold=True, size=7.4, fill="F8FAFC")
    _docx_write_cell(equipment.cell(0, 1), f"{_upper(ctx.get('equipo') or 'EQUIPO')}\n{equipo_meta}", bold=True, size=9.0, fill="FFFFFF")

    tasks_table = document.add_table(rows=2, cols=3)
    _docx_apply_table_grid(tasks_table, [2.0, 2.8, 13.0])
    for idx, title in enumerate(["Puntos de\nBloqueo", "Modo de\nIntervención", "Listado de tareas aplicable al presente procedimiento"]):
        _docx_write_cell(tasks_table.cell(0, idx), title, bold=True, size=7.2, fill="69C97F")
    _docx_write_cell(tasks_table.cell(1, 0), "0", bold=True, size=32.0, fill="FFFFFF")
    _docx_write_cell(tasks_table.cell(1, 1), "MODO 0", bold=True, size=13.0, fill="69C97F")
    tasks_text = "\n".join(f"- {task}" for task in (ctx.get("tareas") or []))
    _docx_write_cell(tasks_table.cell(1, 2), tasks_text, size=6.8, fill="FFFFFF", align="left")

    bar = document.add_table(rows=1, cols=1)
    _docx_apply_table_grid(bar, [17.8])
    _docx_write_cell(bar.cell(0, 0), "Procedimiento - Control de Energías Peligrosas", bold=True, size=7.6, fill="69C97F")

    photo_table = document.add_table(rows=1, cols=1)
    _docx_apply_table_grid(photo_table, [17.8])
    photo_cell = photo_table.cell(0, 0)
    _docx_write_cell(photo_cell, "FOTO DEL PASO A PASO", bold=True, size=11.0, color="94A3B8", fill="FFFFFF")
    photo_bytes, _ = _data_uri_to_bytes(ctx.get("photo_uri", ""))
    if photo_bytes:
        try:
            _docx_clear_cell(photo_cell)
            p = photo_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(io.BytesIO(photo_bytes), width=Cm(12.0))
            _docx_set_cell_borders(photo_cell)
        except Exception:
            _docx_write_cell(photo_cell, "FOTO DEL PASO A PASO", bold=True, size=11.0, color="94A3B8", fill="FFFFFF")

    proc = document.add_table(rows=2, cols=7)
    _docx_apply_table_grid(proc, [1.7, 2.3, 2.0, 2.5, 3.3, 3.8, 2.2])
    proc_headers = ["Punto de\nBloqueo", "Fuente de Energía", "Magnitud", "Ubicación", "Acción", "Verificación", "Dispositivo(s) de\nAislamiento"]
    for idx, title in enumerate(proc_headers):
        _docx_write_cell(proc.cell(0, idx), title, bold=True, size=6.2, color="FFFFFF", fill="3B3B3B")
    proc_values = [
        "No aplica",
        "No Aplica",
        "No Aplica",
        "No Aplica",
        "1. Operar la máquina exclusivamente mediante el panel de Interfaz Hombre-Máquina (HMI) o mediante los controles normales asignados al operador, sin realizar intervención directa sobre el equipo.\n\n2. Mantener la operación dentro de las condiciones normales previstas, sin introducir manos u otros elementos en zonas de riesgo, partes móviles o sectores energizados.\n\n3. No retirar, anular, puentear ni desactivar resguardos, enclavamientos o cualquier dispositivo de seguridad durante la ejecución de la tarea.\n\n4. Detener la actividad y escalar la evaluación si la tarea requiere retirar, superar o intervenir cualquier sistema de protección de la máquina.",
        "1. Confirmar visualmente que todos los resguardos y dispositivos de protección se encuentran instalados, en posición correcta y en condiciones funcionales.\n\n2. Verificar que el equipo opera dentro de sus parámetros normales y que no existe interacción directa del operador con partes móviles, energizadas o zonas de peligro.\n\n3. Comprobar que ninguna protección ha sido modificada, anulada o superada para realizar la actividad.\n\n4. Confirmar que, en caso de requerirse acceso a zonas protegidas o intervención directa sobre el equipo, se haya reevaluado el modo de intervención aplicable, dejando sin efecto el presente procedimiento Modo 0.",
        "No requerido",
    ]
    for idx, value in enumerate(proc_values):
        _docx_write_cell(proc.cell(1, idx), value, size=5.5 if idx in (4, 5) else 6.6, fill="FFFFFF")

    legend_title = document.add_table(rows=1, cols=1)
    _docx_apply_table_grid(legend_title, [17.8])
    _docx_write_cell(legend_title.cell(0, 0), "Clasificación de Energías Peligrosas", bold=True, size=7.0, fill="69C97F")

    energy = document.add_table(rows=2, cols=6)
    _docx_apply_table_grid(energy, [17.8 / 6] * 6)
    energy_items = [
        (0, 0, "E: Eléctrica", "000000", "FFFFFF"),
        (0, 1, "N: Neumática", "0284C7", "FFFFFF"),
        (0, 2, "AM: Amoníaco\nbase gris + franjas naranjas", "D9D9D9", "0F172A"),
        (0, 3, "T: Térmica", "DC2626", "FFFFFF"),
        (0, 4, "H: Hidráulica", "7C3AED", "FFFFFF"),
        (0, 5, "P: Potencial\nbase amarilla + franjas negras", "FFF200", "0F172A"),
        (1, 0, "Q: Química", "FFF200", "0F172A"),
        (1, 1, "V: Vapor", "F59E0B", "111827"),
        (1, 2, "A: Agua", "16A34A", "FFFFFF"),
        (1, 3, "SC: Soda Cáustica\nbase gris + franjas naranjas", "D9D9D9", "0F172A"),
        (1, 4, "Oz: Ozono", "BAE6FD", "0F172A"),
        (1, 5, "GC: Gas Carbónico", "C7D2FE", "111827"),
    ]
    for row, col, label, fill, font_color in energy_items:
        _docx_write_cell(energy.cell(row, col), label, bold=True, size=5.8, color=font_color, fill=fill)

    lock_title = document.add_table(rows=1, cols=1)
    _docx_apply_table_grid(lock_title, [17.8])
    _docx_write_cell(lock_title.cell(0, 0), "Clasificación de Candados según sector y función", bold=True, size=7.0, fill="69C97F")

    locks = document.add_table(rows=1, cols=5)
    _docx_apply_table_grid(locks, [17.8 / 5] * 5)
    lock_items = [
        ("▢\nMMTO\nIndustrial", "EF0000", "FFFFFF"),
        ("▢\nCalidad", "FFF200", "0F172A"),
        ("▢\nProducción", "16A34A", "FFFFFF"),
        ("▢\nMantenimiento Edilicio\nContratistas", "0F7DBD", "FFFFFF"),
        ("▢\nSupervisor MMTO Industrial\n(Bloqueo Departamental)", "050505", "FFFFFF"),
    ]
    for idx, (label, fill, font_color) in enumerate(lock_items):
        _docx_write_cell(locks.cell(0, idx), label, bold=True, size=5.8, color=font_color, fill=fill)

    footer = document.add_table(rows=1, cols=2)
    _docx_apply_table_grid(footer, [8.9, 8.9])
    _docx_write_cell(footer.cell(0, 0), f"Elaborado por: {elaborado_por or '-'}\nPuesto: {puesto_elaborado or '-'} · Fecha: {fecha_firma_txt}", size=6.4, fill="F8FAFC")
    _docx_write_cell(footer.cell(0, 1), f"Aprobado por: {aprobado_por or '-'}\nPuesto: {puesto_aprobado or '-'} · Fecha: {fecha_firma_txt}", size=6.4, fill="F8FAFC")

    for paragraph in document.paragraphs:
        paragraph.paragraph_format.space_before = 0
        paragraph.paragraph_format.space_after = 0

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _xlsx_style_range(ws, cell_range: str, fill: str | None = None, font_color: str = "0F172A", bold: bool = False, size: float = 8.0, align: str = "center") -> None:
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

    side = Side(style="thin", color="111827")
    border = Border(left=side, right=side, top=side, bottom=side)
    for row in ws[cell_range]:
        for cell in row:
            cell.border = border
            cell.alignment = Alignment(horizontal=align, vertical="center", wrap_text=True)
            cell.font = Font(name="Bahnschrift", size=size, bold=bold, color=font_color)
            if fill:
                cell.fill = PatternFill("solid", fgColor=fill.replace("#", ""))


def _xlsx_merge_write(ws, cell_range: str, value: Any, *, fill: str | None = None, font_color: str = "0F172A", bold: bool = False, size: float = 8.0, align: str = "center") -> None:
    top_left = cell_range.split(":", 1)[0]
    ws.merge_cells(cell_range)
    ws[top_left] = value
    _xlsx_style_range(ws, cell_range, fill=fill, font_color=font_color, bold=bold, size=size, align=align)


def _xlsx_add_image(ws, data_uri: str, anchor: str, *, max_width_px: int, max_height_px: int) -> None:
    if not data_uri:
        return
    try:
        from openpyxl.drawing.image import Image as XLImage
        from PIL import Image as PILImage

        image_bytes, _ = _data_uri_to_bytes(data_uri)
        if not image_bytes:
            return
        pil_img = PILImage.open(io.BytesIO(image_bytes)).convert("RGBA")
        pil_img.thumbnail((max_width_px, max_height_px), PILImage.LANCZOS)
        stream = io.BytesIO()
        pil_img.save(stream, format="PNG")
        stream.seek(0)
        xl_img = XLImage(stream)
        xl_img.anchor = anchor
        ws.add_image(xl_img)
    except Exception:
        return


def _xlsx_energy_block(ws, row: int, start_col: int, label: str, fills: list[str], font_color: str = "0F172A") -> None:
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

    side = Side(style="thin", color="111827")
    border = Border(left=side, right=side, top=side, bottom=side)
    for offset, fill in enumerate(fills):
        cell = ws.cell(row=row, column=start_col + offset)
        cell.fill = PatternFill("solid", fgColor=fill)
        cell.border = border
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.font = Font(name="Bahnschrift", size=7, bold=True, color=font_color)
    ws.cell(row=row, column=start_col).value = label


def build_modo_0_excel_bytes(
    ctx: Dict[str, Any],
    *,
    codigo: str,
    revision: str,
    fecha: datetime.date | str,
    organizacion: str,
    logo_uri: str = "",
    personal_afectado: str,
    personal_autorizado: str,
    elaborado_por: str,
    aprobado_por: str,
    puesto_elaborado: str = "",
    puesto_aprobado: str = "",
    fecha_firma: datetime.date | str | None = None,
) -> bytes:
    """Genera un XLSX editable con celdas, textos, estilos, colores y bordes reales."""
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    fecha_txt = _fecha_export_txt(fecha)
    fecha_firma_txt = _fecha_export_txt(fecha_firma) or fecha_txt
    equipo_meta = _equipment_meta(ctx)

    wb = Workbook()
    ws = wb.active
    ws.title = "Modo 0"
    ws.sheet_view.showGridLines = False
    ws.page_setup.orientation = "portrait"
    ws.page_setup.paperSize = ws.PAPERSIZE_A4
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 1
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.page_margins.left = 0.18
    ws.page_margins.right = 0.18
    ws.page_margins.top = 0.18
    ws.page_margins.bottom = 0.18
    ws.page_margins.header = 0
    ws.page_margins.footer = 0

    for col in range(1, 25):
        ws.column_dimensions[get_column_letter(col)].width = 4.2
    for row in range(1, 50):
        ws.row_dimensions[row].height = 18

    white = PatternFill("solid", fgColor="FFFFFF")
    for row in range(1, 50):
        for col in range(1, 25):
            ws.cell(row=row, column=col).fill = white
            ws.cell(row=row, column=col).font = Font(name="Bahnschrift", size=8)

    _xlsx_merge_write(ws, "A1:C3", "" if logo_uri else organizacion, fill="FFFFFF", font_color="B91C1C", bold=True, size=8)
    _xlsx_add_image(ws, logo_uri, "A1", max_width_px=78, max_height_px=52)
    _xlsx_merge_write(ws, "D1:P1", "CONTROL DE ENERGÍAS PELIGROSAS", fill="050505", font_color="FFFFFF", bold=True, size=9)
    _xlsx_merge_write(ws, "D2:P3", "PROCEDIMIENTO ESPECÍFICO PARA CONTROL DE ENERGÍAS", fill="FFFFFF", bold=True, size=10)
    for rng, label, value in (("Q1:S1", "Código", codigo), ("Q2:S2", "Revisión", revision), ("Q3:S3", "Fecha", fecha_txt)):
        _xlsx_merge_write(ws, rng, label, fill="E5E7EB", bold=True, size=8)
        value_rng = rng.replace("Q", "T").replace("S", "X")
        _xlsx_merge_write(ws, value_rng, value, fill="F8FAFC", bold=True, size=8)

    _xlsx_merge_write(ws, "A4:C6", f"Negocio:\n{ctx.get('negocio') or '-'}", fill="FFFFFF", size=8)
    _xlsx_merge_write(ws, "D4:F6", f"Sitio:\n{ctx.get('sitio') or '-'}", fill="FFFFFF", size=8)
    _xlsx_merge_write(ws, "G4:I6", f"Área:\n{ctx.get('area') or '-'}", fill="FFFFFF", size=8)
    _xlsx_merge_write(ws, "J4:M6", f"Línea:\n{ctx.get('linea') or '-'}", fill="FFFFFF", size=8)
    _xlsx_merge_write(ws, "N4:X4", "Personal afectado - Puestos de trabajo", fill="69C97F", bold=True, size=7)
    _xlsx_merge_write(ws, "N5:X5", personal_afectado, fill="FFFFFF", size=7)
    _xlsx_merge_write(ws, "N6:X6", "Personal autorizado - Puestos de trabajo", fill="69C97F", bold=True, size=7)
    _xlsx_merge_write(ws, "N7:X7", personal_autorizado, fill="FFFFFF", size=7)

    _xlsx_merge_write(ws, "A8:C9", "Equipo:", fill="F8FAFC", bold=True, size=8)
    _xlsx_merge_write(ws, "D8:X9", f"{_upper(ctx.get('equipo') or 'EQUIPO')}\n{equipo_meta}", fill="FFFFFF", bold=True, size=11)

    _xlsx_merge_write(ws, "A10:C10", "Puntos de\nBloqueo", fill="69C97F", bold=True, size=7)
    _xlsx_merge_write(ws, "D10:F10", "Modo de\nIntervención", fill="69C97F", bold=True, size=7)
    _xlsx_merge_write(ws, "G10:X10", "Listado de tareas aplicable al presente procedimiento", fill="69C97F", bold=True, size=7)
    _xlsx_merge_write(ws, "A11:C14", "0", fill="FFFFFF", bold=True, size=32)
    _xlsx_merge_write(ws, "D11:F14", "MODO 0", fill="69C97F", bold=True, size=14)
    _xlsx_merge_write(ws, "G11:X14", "\n".join(f"- {task}" for task in (ctx.get("tareas") or [])), fill="FFFFFF", size=7, align="left")

    _xlsx_merge_write(ws, "A15:X15", "Procedimiento - Control de Energías Peligrosas", fill="69C97F", bold=True, size=8)
    _xlsx_merge_write(ws, "A16:X27", "FOTO DEL PASO A PASO", fill="FFFFFF", font_color="94A3B8", bold=True, size=14)
    _xlsx_add_image(ws, ctx.get("photo_uri", ""), "H16", max_width_px=520, max_height_px=210)

    headers = [
        ("A28:C28", "Punto de\nBloqueo"), ("D28:F28", "Fuente de Energía"), ("G28:I28", "Magnitud"),
        ("J28:L28", "Ubicación"), ("M28:P28", "Acción"), ("Q28:U28", "Verificación"), ("V28:X28", "Dispositivo(s) de\nAislamiento"),
    ]
    for rng, label in headers:
        _xlsx_merge_write(ws, rng, label, fill="3B3B3B", font_color="FFFFFF", bold=True, size=7)
    _xlsx_merge_write(ws, "A29:C35", "No aplica", fill="FFFFFF", size=7)
    _xlsx_merge_write(ws, "D29:F35", "No Aplica", fill="FFFFFF", size=7)
    _xlsx_merge_write(ws, "G29:I35", "No Aplica", fill="FFFFFF", size=7)
    _xlsx_merge_write(ws, "J29:L35", "No Aplica", fill="FFFFFF", size=7)
    accion = "1. Operar la máquina exclusivamente mediante el panel de Interfaz Hombre-Máquina (HMI) o mediante los controles normales asignados al operador, sin realizar intervención directa sobre el equipo.\n\n2. Mantener la operación dentro de las condiciones normales previstas, sin introducir manos u otros elementos en zonas de riesgo, partes móviles o sectores energizados.\n\n3. No retirar, anular, puentear ni desactivar resguardos, enclavamientos o cualquier dispositivo de seguridad durante la ejecución de la tarea.\n\n4. Detener la actividad y escalar la evaluación si la tarea requiere retirar, superar o intervenir cualquier sistema de protección de la máquina."
    verificacion = "1. Confirmar visualmente que todos los resguardos y dispositivos de protección se encuentran instalados, en posición correcta y en condiciones funcionales.\n\n2. Verificar que el equipo opera dentro de sus parámetros normales y que no existe interacción directa del operador con partes móviles, energizadas o zonas de peligro.\n\n3. Comprobar que ninguna protección ha sido modificada, anulada o superada para realizar la actividad.\n\n4. Confirmar que, en caso de requerirse acceso a zonas protegidas o intervención directa sobre el equipo, se haya reevaluado el modo de intervención aplicable, dejando sin efecto el presente procedimiento Modo 0."
    _xlsx_merge_write(ws, "M29:P35", accion, fill="FFFFFF", size=6)
    _xlsx_merge_write(ws, "Q29:U35", verificacion, fill="FFFFFF", size=6)
    _xlsx_merge_write(ws, "V29:X35", "No requerido", fill="FFFFFF", size=7)

    _xlsx_merge_write(ws, "A36:X36", "Clasificación de Energías Peligrosas", fill="69C97F", bold=True, size=8)
    energy_blocks = [
        (37, 1, "E: Eléctrica", ["000000", "000000", "000000", "000000"], "FFFFFF"),
        (37, 5, "N: Neumática", ["0284C7", "0284C7", "0284C7", "0284C7"], "FFFFFF"),
        (37, 9, "AM: Amoníaco", ["D9D9D9", "F59E0B", "D9D9D9", "F59E0B"], "0F172A"),
        (37, 13, "T: Térmica", ["DC2626", "DC2626", "DC2626", "DC2626"], "FFFFFF"),
        (37, 17, "H: Hidráulica", ["7C3AED", "7C3AED", "7C3AED", "7C3AED"], "FFFFFF"),
        (37, 21, "P: Potencial", ["FFF200", "050505", "FFF200", "050505"], "0F172A"),
        (38, 1, "Q: Química", ["FFF200", "FFF200", "FFF200", "FFF200"], "0F172A"),
        (38, 5, "V: Vapor", ["F59E0B", "F59E0B", "F59E0B", "F59E0B"], "111827"),
        (38, 9, "A: Agua", ["16A34A", "16A34A", "16A34A", "16A34A"], "FFFFFF"),
        (38, 13, "SC: Soda Cáustica", ["D9D9D9", "F59E0B", "D9D9D9", "F59E0B"], "0F172A"),
        (38, 17, "Oz: Ozono", ["BAE6FD", "BAE6FD", "BAE6FD", "BAE6FD"], "0F172A"),
        (38, 21, "GC: Gas Carbónico", ["C7D2FE", "C7D2FE", "C7D2FE", "C7D2FE"], "111827"),
    ]
    for item in energy_blocks:
        _xlsx_energy_block(ws, *item)

    _xlsx_merge_write(ws, "A39:X39", "Clasificación de Candados según sector y función", fill="69C97F", bold=True, size=8)
    locks = [
        ("A40:E41", "▢  MMTO\nIndustrial", "EF0000", "FFFFFF"),
        ("F40:J41", "▢  Calidad", "FFF200", "0F172A"),
        ("K40:N41", "▢  Producción", "16A34A", "FFFFFF"),
        ("O40:S41", "▢  Mantenimiento Edilicio\nContratistas", "0F7DBD", "FFFFFF"),
        ("T40:X41", "▢  Supervisor MMTO Industrial\n(Bloqueo Departamental)", "050505", "FFFFFF"),
    ]
    for rng, label, fill, font_color in locks:
        _xlsx_merge_write(ws, rng, label, fill=fill, font_color=font_color, bold=True, size=7)

    _xlsx_merge_write(ws, "A42:L44", f"Elaborado por: {elaborado_por or '-'}\nPuesto: {puesto_elaborado or '-'} · Fecha: {fecha_firma_txt}", fill="F8FAFC", size=7)
    _xlsx_merge_write(ws, "M42:X44", f"Aprobado por: {aprobado_por or '-'}\nPuesto: {puesto_aprobado or '-'} · Fecha: {fecha_firma_txt}", fill="F8FAFC", size=7)

    for row in range(37, 39):
        ws.row_dimensions[row].height = 26
    ws.row_dimensions[40].height = 26
    ws.row_dimensions[41].height = 26

    ws.print_area = "A1:X44"
    output = io.BytesIO()
    wb.save(output)
    return output.getvalue()



# ============================================================
# INTERFAZ STREAMLIT
# ============================================================

st.markdown(
    """
    <style>
        html, body, [class*="css"], .stApp, input, textarea, button {
            font-family: Bahnschrift, 'Bahnschrift SemiCondensed', 'Arial Narrow', Arial, Helvetica, sans-serif !important;
        }
    </style>
    """,
    unsafe_allow_html=True,
)

st.title("GENERADOR DE PROCEDIMIENTOS MODO 0")
st.markdown(
    """
    <div class="soft-note">
        Importá el borrador JSON generado por la app de análisis. Esta primera versión genera el procedimiento específico
        para <strong>Modo 0</strong>, manteniendo una estética tipo ISO y estructura documental lista para descarga en Word DOCX y Excel XLSX.
    </div>
    """,
    unsafe_allow_html=True,
)

with st.sidebar:
    st.subheader("Configuración documental")
    codigo = st.text_input("Código", value="LOTO-M0-001")
    revision = st.text_input("Revisión", value="00")
    fecha = st.date_input("Fecha", value=datetime.date.today(), format="DD/MM/YYYY")
    organizacion = st.text_input("Organización / logo textual", value="ARCA CONTINENTAL")
    logo_file = st.file_uploader(
        "Logo opcional para reemplazar arca.png",
        type=["png", "jpg", "jpeg", "webp"],
        help="Si no se carga un archivo, el documento intenta usar arca.png desde la carpeta de la aplicación.",
    )

    st.divider()
    st.subheader("Firmas")
    elaborado_por = st.text_input("Elaborado por", value="")
    puesto_elaborado = st.text_input("Puesto de quien elabora", value="")
    aprobado_por = st.text_input("Aprobado por", value="")
    puesto_aprobado = st.text_input("Puesto de quien aprueba", value="")
    fecha_firma = st.date_input("Fecha de firmas", value=fecha, format="DD/MM/YYYY")

uploaded_json = st.file_uploader(
    "Importar borrador JSON",
    type=["json"],
    help="Usá el archivo generado por el botón 'Descargar borrador JSON' de la app anterior.",
)

if uploaded_json is None:
    st.info("Cargá un borrador JSON para generar el procedimiento.")
    st.stop()

try:
    payload = json.loads(uploaded_json.getvalue().decode("utf-8"))
except Exception as exc:
    st.error(f"No se pudo leer el JSON importado: {exc}")
    st.stop()

ctx_detected = extract_procedure_context(payload)

st.subheader("Datos importados editables")
st.caption("Todos los campos detectados desde el JSON pueden reescribirse antes de generar el procedimiento.")

id_col_1, id_col_2, id_col_3 = st.columns(3)
with id_col_1:
    negocio_edit = st.text_input("Negocio", value=_normalize(ctx_detected.get("negocio")), key="edit_negocio")
    sitio_edit = st.text_input("Sitio", value=_normalize(ctx_detected.get("sitio")), key="edit_sitio")
    area_edit = st.text_input("Área", value=_normalize(ctx_detected.get("area")), key="edit_area")
with id_col_2:
    linea_edit = st.text_input("Línea", value=_normalize(ctx_detected.get("linea")), key="edit_linea")
    equipo_edit = st.text_input("Equipo", value=_normalize(ctx_detected.get("equipo")), key="edit_equipo")
    fabricante_edit = st.text_input("Fabricante", value=_normalize(ctx_detected.get("fabricante")), key="edit_fabricante")
with id_col_3:
    modelo_edit = st.text_input("Modelo", value=_normalize(ctx_detected.get("modelo")), key="edit_modelo")
    anio_edit = st.text_input("Año", value=_normalize(ctx_detected.get("anio")), key="edit_anio")
    modo_inicial_edit = st.text_input("Modo inicial", value=_normalize(ctx_detected.get("modo_inicial")), key="edit_modo_inicial")

modo_final_edit = st.text_input("Modo final", value=_normalize(ctx_detected.get("modo_final")), key="edit_modo_final")

st.subheader("Personal editable")
st.caption("Estos campos se escriben manualmente antes de generar el procedimiento y se imprimen en el encabezado del documento.")
person_col_1, person_col_2 = st.columns(2)
with person_col_1:
    personal_afectado = st.text_area(
        "Personal afectado - Puestos de trabajo",
        value="Técnicos del Sector de Mantenimiento.\nContratistas.",
        height=96,
        key="edit_personal_afectado",
    )
with person_col_2:
    personal_autorizado = st.text_area(
        "Personal autorizado - Puestos de trabajo",
        value="Técnicos del Sector Mantenimiento.\nContratistas.",
        height=96,
        key="edit_personal_autorizado",
    )

st.subheader("Tareas aplicables")
tasks_detected = ctx_detected.get("tareas") or []
tasks_text = st.text_area(
    "Listado de tareas desde tareas_predefinidas",
    value="\n".join(tasks_detected),
    height=120,
    help="Una tarea por línea. El valor inicial se toma prioritariamente desde la clave tareas_predefinidas del JSON.",
)

st.subheader("Foto del paso a paso")
step_photo_file = st.file_uploader(
    "Subir foto o imagen del paso a paso",
    type=["png", "jpg", "jpeg", "webp"],
    help="Esta imagen reemplaza la foto del equipo del JSON y se inserta en el área central del procedimiento.",
)
photo_uri = _uploaded_file_data_uri(step_photo_file)
if step_photo_file is not None:
    st.image(step_photo_file, caption="Foto del paso a paso cargada", use_container_width=True)
else:
    st.caption("Si no se sube una imagen, el procedimiento mostrará un recuadro reservado para la foto del paso a paso.")

with st.expander("Lienzo opcional para tags sobre imagen"):
    st.caption("El lienzo se habilita solo si el entorno tiene instalada la extensión streamlit-drawable-canvas. En caso contrario, se conserva la imagen cargada sin anotaciones.")
    try:
        from streamlit_drawable_canvas import st_canvas  # type: ignore
        from PIL import Image  # type: ignore
        if step_photo_file is None:
            st.info("Subí primero una foto del paso a paso para activar el lienzo de marcado.")
        else:
            canvas_result = st_canvas(
                fill_color="rgba(255, 0, 0, 0.15)",
                stroke_width=3,
                stroke_color="#C00000",
                background_image=Image.open(step_photo_file),
                update_streamlit=True,
                height=360,
                drawing_mode="freedraw",
                key="step_photo_canvas",
            )
            if canvas_result.image_data is not None:
                annotated = Image.fromarray(canvas_result.image_data.astype("uint8"), "RGBA")
                buffer = __import__("io").BytesIO()
                annotated.save(buffer, format="PNG")
                photo_uri = _bytes_to_data_uri(buffer.getvalue(), "image/png")
                st.success("Se usará la imagen anotada del lienzo en el procedimiento.")
    except Exception:
        st.info("Lienzo no disponible en este entorno. Podés subir la foto ya marcada o continuar con la imagen sin anotaciones.")

ctx = dict(ctx_detected)
ctx.update(
    {
        "negocio": negocio_edit,
        "sitio": sitio_edit,
        "area": area_edit,
        "linea": linea_edit,
        "equipo": equipo_edit,
        "fabricante": fabricante_edit,
        "modelo": modelo_edit,
        "anio": anio_edit,
        "modo_inicial": modo_inicial_edit,
        "modo_final": modo_final_edit,
        "tareas": _split_tasks(tasks_text),
        "photo_uri": photo_uri,
    }
)
modo_final = ctx.get("modo_final")

col_a, col_b, col_c = st.columns(3)
with col_a:
    st.metric("Modo final", _normalize(modo_final) or "Sin dato")
with col_b:
    st.metric("Equipo", _normalize(ctx.get("equipo"))[:32] or "Sin dato")
with col_c:
    st.metric("Tareas", len(ctx.get("tareas") or []))

if not _mode_is_modo_0(modo_final):
    st.markdown(
        f"""
        <div class="status-warn">
            El borrador importado no corresponde a <strong>Modo 0</strong>. Modo final detectado: <strong>{_html(modo_final or 'sin dato')}</strong>.
            Esta versión inicial del generador solo habilita la plantilla Modo 0. Para Modo 1, 2 o 3 se deberá agregar una plantilla específica.
        </div>
        """,
        unsafe_allow_html=True,
    )
    with st.expander("Ver datos detectados del JSON"):
        st.json({"identificacion": ctx, "computed": payload.get("computed", {})})
    st.stop()

logo_uri = _logo_data_uri(logo_file)
procedure_html = build_modo_0_html(
    ctx,
    codigo=codigo,
    revision=revision,
    fecha=fecha,
    organizacion=organizacion,
    logo_uri=logo_uri,
    personal_afectado=personal_afectado,
    personal_autorizado=personal_autorizado,
    elaborado_por=elaborado_por,
    aprobado_por=aprobado_por,
    puesto_elaborado=puesto_elaborado,
    puesto_aprobado=puesto_aprobado,
    fecha_firma=fecha_firma,
)

st.markdown(
    """
    <div class="status-ok">
        Borrador compatible con <strong>Modo 0</strong>. Se generó la vista previa del procedimiento y quedaron habilitadas las descargas.
    </div>
    """,
    unsafe_allow_html=True,
)

st.subheader("Vista previa del procedimiento")
components.html(procedure_html, height=850, scrolling=True)

file_stamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
base_filename = f"Procedimiento_Modo_0_{file_stamp}"

col_word, col_excel, col_html, col_json = st.columns(4)
with col_word:
    st.download_button(
        "Descargar Word DOCX",
        data=html_to_word_bytes(
            ctx,
            codigo=codigo,
            revision=revision,
            fecha=fecha,
            organizacion=organizacion,
            logo_uri=logo_uri,
            personal_afectado=personal_afectado,
            personal_autorizado=personal_autorizado,
            elaborado_por=elaborado_por,
            aprobado_por=aprobado_por,
            puesto_elaborado=puesto_elaborado,
            puesto_aprobado=puesto_aprobado,
            fecha_firma=fecha_firma,
        ),
        file_name=f"{base_filename}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )

with col_excel:
    try:
        excel_bytes = build_modo_0_excel_bytes(
            ctx,
            codigo=codigo,
            revision=revision,
            fecha=fecha,
            organizacion=organizacion,
            logo_uri=logo_uri,
            personal_afectado=personal_afectado,
            personal_autorizado=personal_autorizado,
            elaborado_por=elaborado_por,
            aprobado_por=aprobado_por,
            puesto_elaborado=puesto_elaborado,
            puesto_aprobado=puesto_aprobado,
            fecha_firma=fecha_firma,
        )
        st.download_button(
            "Descargar Excel XLSX",
            data=excel_bytes,
            file_name=f"{base_filename}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
        )
    except Exception as exc:
        st.warning(f"Excel editable no disponible en este entorno. Detalle técnico: {exc}")

with col_html:
    st.download_button(
        "Descargar HTML",
        data=procedure_html.encode("utf-8"),
        file_name=f"{base_filename}.html",
        mime="text/html",
        use_container_width=True,
    )

with col_json:
    st.download_button(
        "Descargar JSON importado",
        data=json.dumps(payload, ensure_ascii=False, indent=4).encode("utf-8"),
        file_name=f"Borrador_importado_{file_stamp}.json",
        mime="application/json",
        use_container_width=True,
    )

with st.expander("Datos editados para generación"):
    st.json(
        {
            "negocio": ctx.get("negocio"),
            "sitio": ctx.get("sitio"),
            "area": ctx.get("area"),
            "linea": ctx.get("linea"),
            "equipo": ctx.get("equipo"),
            "modo_inicial": ctx.get("modo_inicial"),
            "modo_final": ctx.get("modo_final"),
            "tareas": ctx.get("tareas"),
            "foto_paso_a_paso_cargada": bool(ctx.get("photo_uri")),
        }
    )
